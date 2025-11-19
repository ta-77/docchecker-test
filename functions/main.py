# functions/main.py

import os
import re
import io
import json
from firebase_admin import initialize_app
from firebase_functions import https_fn, options
from docx import Document
from docx.shared import Pt

# --- ★ 修正点 1: 新しいライブラリのインポート ---
# 'google.generativeai' ではなく 'google.genai' から GenerativeModel を直接インポート
from google.genai import GenerativeModel
# --- ★ 修正ここまで ---

# Firebaseアプリの初期化
initialize_app()

# --- ★ APIキー設定（ローカル / 本番 両対応）---
api_key = os.environ.get("GEMINI_API_KEY")

if not api_key:
    api_key = options.SecretParam("GEMINI_API_KEY")

# --- ★ 修正点 2: genai.configure() は不要になったので削除 ---
# (genai.configure(api_key=api_key) は削除)
# --- ★ 修正ここまで ---


# --- ★ここからがバックエンドAPI本体 ---

@https_fn.on_request(
    region="us-central1",  # ローカルエミュレータ用 (テスト用)
    memory=options.MemoryOption.GB_1,
    timeout_sec=300,
    cors=options.CorsOptions(
        cors_origins=[
            "http://127.0.0.1:5000", # ローカルエミュレータ用
            "http://localhost:5000",  # ローカルエミュレータ用
            "https*://doccheck-test.web.app" # 本番のHosting URL
        ],
        cors_methods=["post"]
    )
)
def checkDocument(req: https_fn.Request) -> https_fn.Response:
    """
    Wordファイルを受け取り、ルールベースおよびAIベースのチェックを実行する
    HTTP Function。
    """

    # --- FRB-1.1: リクエスト検証 ---
    if req.method != "POST":
        return https_fn.Response(
            {"detail": f"Method {req.method} not allowed. Use POST."},
            status=405,
            mimetype="application/json"
        )

    if not req.files or 'file' not in req.files:
        return https_fn.Response(
            {"detail": "File not found in request"},
            status=400,
            mimetype="application/json"
        )

    uploaded_file = req.files['file']

    if uploaded_file.content_type != 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        return https_fn.Response(
            {"detail": "Invalid file type. Only .docx is allowed."},
            status=400,
            mimetype="application/json"
        )

    # --- ★ ここからが単一の try ブロック ---
    try:
        # === FRB-2.0: Word文書の解析 ===
        document_structure = []
        full_text_for_ai = [] # AIに渡すためのプレーンテキスト
        
        # 1. ファイルをメモリ上で読み込む
        file_stream = io.BytesIO(uploaded_file.read())
        document = Document(file_stream)

        # === FRB-2.2 & 2.3 & 2.4: ルールベースチェック ===
        
        # ルール定義 (例: 本文の1行目インデントは 10.5pt = 1文字)
        REQUIRED_FIRST_LINE_INDENT = Pt(10.5) 
        
        for para in document.paragraphs:
            paragraph_text = para.text
            full_text_for_ai.append(paragraph_text) # AI用にテキストを収集
            
            paragraph_runs = []
            paragraph_errors = []

            # --- FRB-2.4 (インデントチェック) ---
            if para.style.name == '本文':
                indent = para.paragraph_format.first_line_indent
                if indent is None or indent != REQUIRED_FIRST_LINE_INDENT:
                    paragraph_errors.append({
                        "type": "IndentError",
                        "message": f"本文スタイルのインデントが不正です (現在値: {indent})"
                    })

            for run in para.runs:
                run_text = run.text
                run_errors = []
                
                # 表示/判定用のフォント名を初期化
                run_font_to_check = None

                # --- FRB-2.3 (フォントチェック) ---
                
                # ルールA: 半角数字は Century
                if re.search(r'^[0-9]+$', run_text.strip()):
                    run_font_to_check = run.font.name # ラテン文字フォントを取得
                    if run_font_to_check != 'Century':
                        run_errors.append({
                            "type": "FontError",
                            "message": f"フォントが 'Century' ではありません (現在: {run_font_to_check})"
                        })
                
                # ルールB: 日本語を含む場合は MS明朝
                elif re.search(r'[ぁ-んァ-ヶ一-龠]', run_text):
                    
                    # .east_asia 属性が存在するかどうかを先に確認する (hasattr)
                    if hasattr(run.font, 'east_asia'):
                        run_font_to_check = run.font.east_asia
                        if run_font_to_check != 'MS明朝':
                            run_errors.append({
                                "type": "FontError",
                                "message": f"フォント(東アジア)が 'MS明B' ではありません (現在: {run_font_to_check})"
                            })
                    else:
                        # .east_asia が存在しない場合、.name (ラテン文字) にフォールバックする
                        run_font_to_check = run.font.name
                        if run_font_to_check != 'MS明朝':
                            run_errors.append({
                                "type": "FontError",
                                "message": f"フォント(標準)が 'MS明朝' ではありません (現在: {run_font_to_check})"
                            })
                
                # その他の文字（記号など）
                else:
                    run_font_to_check = run.font.name # とりあえずラテン文字設定を格納
                
                paragraph_runs.append({
                    "text": run_text,
                    "font": str(run_font_to_check), # (Noneの可能性を考慮し str() で囲む)
                    "errors": run_errors
                })

            document_structure.append({
                "type": "paragraph",
                "text": paragraph_text,
                "style": para.style.name,
                "errors": paragraph_errors,
                "runs": paragraph_runs
            })

        # === FRB-3.0: AIによる内容チェック ===
        ai_suggestions = []
        combined_text = "\n".join(full_text_for_ai)
        
        # テキストが空でない場合のみAIチェックを実行
        if combined_text.strip():
            
            # --- ★ 修正点 3: 新しいライブラリ (google-genai) の呼び出し方 ---
            # クライアントの初期化 (api_key は関数の先頭で読み込んだものを使用)
            client = GenerativeModel(
                model_name="gemini-1.5-flash", # ★ ご希望の flash モデル
                api_key=api_key
            )
            # --- ★ 修正ここまで ---
            
            prompt = f"""
            あなたは優秀なビジネス文書の校閲者です。
            以下のテキストを読み、[1] 敬語やビジネス表現の誤り、[2] 論理的な矛盾や不明瞭な点、[3] 必須項目（日付、担当者名など）の欠落の可能性、を指摘してください。
            指摘事項のみを、簡潔なメッセージのリストとしてJSON配列の形式で回答してください。
            
            例:
            [
              {{ "message": "「〜していただく」は二重敬語の可能性があります。" }},
              {{ "message": "日付の記載が見当たりません。" }}
            ]

            テキスト:
            ---
            {combined_text}
            ---
            """

            try:
                # --- ★ 修正点 4: 新しいライブラリでの呼び出し ---
                response = client.generate_content(prompt)
                # --- ★ 修正ここまで ---
                
                # AIの回答をパース ( "```json\n[...]\n```" のようなマークダウンを除去)
                clean_response = re.sub(r'```json\n?|\n?```', '', response.text.strip())
                
                # JSONとして解析
                ai_suggestions = json.loads(clean_response)

            except Exception as ai_e:
                print(f"AI Error: {ai_e}")
                # AIが失敗しても、ルールチェックの結果は返す
                ai_suggestions = [
                    {"message": f"AIによるチェック中にエラーが発生しました: {str(ai_e)}"}
                ]
        else:
            ai_suggestions = [
                {"message": "文書が空のようです。AIチェックはスキップされました。"}
            ]


        # === FRB-4.0: レスポンス合成 ===
        final_result = {
            "documentStructure": document_structure,
            "aiSuggestions": ai_suggestions
        }

        # --- FRB-4.2: レスポンス返却 ---
        return https_fn.Response(
            json.dumps(final_result, ensure_ascii=False), # JSON文字列に変換
            status=200,
            mimetype="application/json"
        )

    except Exception as e:
        # --- FRB-1.3: 内部エラー処理 ---
        print(f"Internal Server Error: {e}") 
        return https_fn.Response(
            {"detail": f"Internal server error: {str(e)}"},
            status=500,
            mimetype="application/json"
        )